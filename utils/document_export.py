from docx import Document

def export_to_docx(questions, course, topics, question_types, difficulty):
    doc = Document()
    doc.add_heading(f'{course} - {topics} - Questions ({difficulty})', 0)

    for i, question in enumerate(questions):
        doc.add_heading(f'Q{i+1} ({question["type"]})', level=1)
        doc.add_paragraph(question['question'])
        if 'answer' in question and question['answer']:
            doc.add_heading('Answer:', level=2)
            doc.add_paragraph(question['answer'])

    doc.save('questions.docx')
